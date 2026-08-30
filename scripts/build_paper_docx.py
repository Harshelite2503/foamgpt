"""Build the FoamGPT draft paper (docx + md) from live corpus/dataset/benchmark results."""
import json
from pathlib import Path
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from foamgpt.config import BENCH_DIR, RAW_DIR, ROOT
from foamgpt.curate.normalize import CURATED

# ---------------- numbers ----------------
rows = [json.loads(l) for l in (RAW_DIR / "papers.jsonl").read_text().splitlines()]
rel = [r for r in rows if r["relevance"] == 2 and r["year"]]
papers = {r["id"]: r for r in rows}
n_oa = sum(1 for r in rel if r["oa_pdf_url"])
ex = [json.loads(l) for l in (ROOT / "data/extracted/extractions.jsonl").read_text().splitlines()]
n_proc = len(ex); n_on = sum(e["extraction"]["is_syntactic_foam_paper"] for e in ex)
df = pd.read_csv(CURATED); df["year"] = df.paper_id.map(lambda i: papers[i]["year"])
prim = int((df.data_origin == "primary").sum()); sec = int((df.data_origin == "secondary").sum()); mod = int((df.data_origin == "model").sum())
flagged = int((df["flags"].fillna("") != "").sum())
ml = pd.read_csv(BENCH_DIR / "ml_baselines.csv")
llm_path = BENCH_DIR / "llm_agent_summary.csv"
llm = pd.read_csv(llm_path) if llm_path.exists() else None
onpaper = df.groupby("paper_id")[["strength_mpa", "modulus_mpa", "measured_density_g_cc"]].apply(lambda g: g.notna().any().any())
yrs = pd.Series({p: papers[p]["year"] for p in onpaper.index})
pre, post = onpaper[yrs < 2010].mean(), onpaper[yrs >= 2010].mean()
clean = df[(df["flags"].fillna("") == "") & (df.data_origin == "primary") & (df.test_type == "compression")]
al = clean[clean.matrix_class == "aluminum"].dropna(subset=["measured_density_g_cc", "strength_mpa"])
r_al = np.corrcoef(al.measured_density_g_cc, al.strength_mpa)[0, 1]
val = pd.read_csv(ROOT / "data/curated/validation_sample_100.csv")

md = []  # markdown mirror
doc = Document(); doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(11)
for s in doc.sections: s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(1)
def H(t, l=1): doc.add_heading(t, level=l); md.append("#" * l + " " + t + "\n")
def P(t, italic=False, align=None):
    p = doc.add_paragraph(); r = p.add_run(t); r.italic = italic
    if align is not None: p.alignment = align
    md.append(t + "\n")
def B(items):
    for i in items: doc.add_paragraph(i, style="List Bullet"); md.append("- " + i)
    md.append("")
def T(header, body, widths=None, caption=None):
    if caption: P(caption, italic=True)
    t = doc.add_table(rows=1, cols=len(header)); t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]; c.text = ""; c.paragraphs[0].add_run(h).bold = True
    for row in body:
        cells = t.add_row().cells
        for i, v in enumerate(row): cells[i].text = str(v)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths): row.cells[i].width = Inches(w)
    doc.add_paragraph()
    md.append("| " + " | ".join(header) + " |"); md.append("|" + "---|" * len(header))
    for row in body: md.append("| " + " | ".join(str(v) for v in row) + " |")
    md.append("")
def FIG(path, cap, w=6.0):
    doc.add_picture(str(path), width=Inches(w)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    P(cap, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER); md.insert(len(md) - 1, f"![]({Path(path).relative_to(ROOT)})")

# ---------------- title ----------------
title = "FoamGPT: A Literature-Scale Process–Structure–Property Dataset for Syntactic Foams Extracted with Large Language Models, and What Machine Learning and LLMs Can Predict From It"
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run(title); r.bold = True; r.font.size = Pt(16)
md.append("# " + title + "\n")
P("Harsh Vardhan Gupta¹, Nikhil Gupta¹\n¹ Department of Mechanical and Aerospace Engineering, NYU Tandon School of Engineering, Brooklyn, NY, USA", align=WD_ALIGN_PARAGRAPH.CENTER)
P("DRAFT v0.1 — August 2026 — for internal discussion", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

# ---------------- abstract ----------------
H("Abstract")
llm_sentence = ""
if llm is not None and len(llm):
    s = llm.set_index(["target", "condition"])
    try:
        zs = s.loc[("strength_mpa", "zero_shot")]; rg = s.loc[("strength_mpa", "rag_5")]
        llm_sentence = (f" A frontier LLM asked to predict compressive strength from the recipe alone achieves a median absolute "
                        f"percentage error of {zs.median_ape_pct:.0f}% with {zs.coverage_90:.0%} empirical coverage of its stated 90% intervals; "
                        f"supplying five retrieved records from other papers changes these to {rg.median_ape_pct:.0f}% and {rg.coverage_90:.0%}.")
    except KeyError:
        pass
P(f"Syntactic foams — hollow-particle-filled composites — have a three-decade experimental literature whose quantitative results remain locked in PDF tables and figures. We present FoamGPT, an open pipeline and dataset that converts that literature into a machine-readable process–structure–property (PSP) table with per-value provenance. From {len(rel):,} relevant papers identified in OpenAlex (1990–2026) we processed the {n_proc} open-access papers, of which {n_on} were confirmed on-topic, and extracted {len(df)} records ({prim} primary measurements, {sec} literature-quoted values, {mod} model results) using a large language model constrained by a 40-field schema, a no-guessing rule, and mandatory verbatim evidence quotes. Automated physical-range checks flag {flagged} records for expert review. The extracted data reproduce known structure–property trends without post-processing. Using paper-level cross-validation, a random forest predicts foam density from processing descriptors alone (R² = {ml[ml.target=='measured_density_g_cc'].r2.max():.2f}) but cannot yet predict modulus or strength across laboratories (R² < 0), quantifying how far recipe-level features are from transferable mechanical prediction at this data scale.{llm_sentence} We release the dataset, extraction prompts, and benchmark code, and describe a validation protocol and the remaining barriers — figure-only reporting, paywalled corpora, and unit/definition heterogeneity — to a complete literature-scale dataset.")

# ---------------- 1 Introduction ----------------
H("1. Introduction")
P("Syntactic foams are particulate composites in which hollow microspheres (glass, ceramic, fly-ash cenosphere, polymer, or metal) are dispersed in a polymer or metal matrix. Their combination of low density, high specific compressive strength, damage tolerance, and low moisture uptake makes them the material of choice for deep-submergence buoyancy, marine and aerospace sandwich cores, and energy-absorbing structures. Because properties depend on the full processing chain — matrix chemistry, microsphere grade and wall-thickness ratio, volume fraction, mixing and curing route — the field has produced a large empirical literature: we identify " + f"{len(rel):,} relevant papers since 1990, growing every decade.")
P("Almost none of this knowledge is usable computationally. Each paper reports its own compositions in its own tables, with inconsistent units (MPa vs. GPa; kg/m³ vs. g/cm³; vol% vs. wt%), paper-specific sample labels, and results that are often shown only graphically. Consequently, materials selection is performed by manual literature review, and machine-learning studies in the field are trained on a few dozen in-house samples that do not generalise across laboratories.")
P("Large language models (LLMs) now extract structured data from scientific text with near-expert accuracy when the schema is explicit and outputs are validated [Schilling-Wilhelmi 2025; Zimmermann 2025]. Nearly all demonstrations target chemistry and crystalline materials; mechanical and processing data for composites is a recognised gap [Van 2026]. At the same time, LLMs are increasingly consulted for materials questions, and whether their numerical answers reflect knowledge or plausible fabrication has not been measured for composite properties [Zhang 2026].")
P("This paper makes three contributions. (1) An open, provenance-carrying PSP dataset for syntactic foams built from the open-access literature with an LLM extraction pipeline that refuses to guess. (2) A quantitative characterisation of what the literature actually contains — how much is tabulated versus figure-only, how coverage varies by matrix class and era, and what internal inconsistencies occur. (3) A benchmark of classical ML and a frontier LLM on cross-laboratory property prediction, with calibration analysis for the LLM.")

# ---------------- 2 Related work ----------------
H("2. Related Work")
P("Syntactic foam mechanics. The roles of microballoon wall-thickness ratio, volume fraction, matrix modulus, strain rate and temperature have been mapped in polymer-matrix foams and, with pressure-infiltration processing, in aluminium and magnesium matrix foams [Gupta & Rohatgi; Orbulov; Szlancsik]. Micromechanical models (Bardella–Genna, Porfiri–Gupta) predict elastic moduli from constituent properties, but strength and energy absorption remain empirical.")
P("LLM-based scientific data extraction. Structured extraction with schema constraints and validation has been demonstrated for reaction conditions, MOF synthesis, and polymer properties. Reviews [Jiang 2025; Van 2026] identify composites/manufacturing as under-served and highlight provenance and hallucination control as open problems.")
P("Benchmarks of LLM materials knowledge. MatSciBench [Zhang 2026] evaluates reasoning on textbook-style problems; van Herck et al. [2025] compare fine-tuned LLMs with ML on tabular property prediction. No prior work measures calibration of LLM numerical predictions for a composite class against literature ground truth.")

# ---------------- 3 Methods ----------------
H("3. Methods")
H("3.1 Corpus construction", 2)
P(f"Seven queries (\"syntactic foam(s)\", \"hollow glass microspheres\", \"glass microballoons\", \"cenosphere composite\", \"hollow particle composite\", \"metal matrix syntactic foam\") were run against the OpenAlex works index restricted to 1990–2026, returning 4,912 works. A title/abstract relevance score labelled {len(rel):,} as strongly relevant (syntactic foam or microballoon explicitly named) and 662 as weakly relevant. Only {n_oa} of the strongly relevant papers ({n_oa/len(rel):.0%}) expose an open-access PDF; {n_proc} of those resolved to a real PDF (the remainder are HTML landing pages) and form the pilot corpus. Figure 1 shows the corpus by year.")
FIG(BENCH_DIR / "corpus_by_year.png", "Figure 1. Strongly relevant syntactic-foam papers per year in the OpenAlex-derived corpus.")
H("3.2 Parsing", 2)
P("PDFs were converted with PyMuPDF; detected tables were rendered as Markdown so that rows and columns survive. Full-length documents were retained (an initial 40-page cap silently truncated ten theses and was removed; the truncated papers were re-extracted). Image-only scans (one in the corpus) and bitmap tables yield only captions and are recorded as such.")
H("3.3 Schema", 2)
P("One record is one composition tested under one condition. Fields are grouped as Processing (matrix class and name; particle type, grade, true density, mean diameter, wall-thickness ratio η; volume and weight fraction; process route; cure temperature; additional fillers), Structure (measured and theoretical density, matrix porosity, particle breakage), Test (type, strain rate, temperature, standard), and Properties (modulus, strength, failure strain, energy absorption, plateau stress, densification strain, DMA storage modulus and tan δ, thermal conductivity, CTE, moisture uptake). Every numeric field is optional. Each record carries a data-origin tag (primary measurement, secondary literature quotation with citation, or model result), at least one verbatim evidence quote with location, and an extractor confidence.")
H("3.4 Extraction", 2)
P("Extraction is performed by Claude (Anthropic) with the schema inlined in the system prompt and a fixed rule set: one record per composition × condition; never invent a value — leave it null if it is not stated numerically; convert to schema units; do not read values off plots; tag literature-quoted values as secondary. Outputs are validated against the Pydantic schema with one self-repair turn, then passed through range and unit checks. Two execution lanes share the same prompt, schema and validator: an API lane (Claude Opus 5 via the Messages API; three papers, ≈US$1) and an agent lane in which Claude Code subagents (Claude Opus 5) read the paper text and write the JSON, used for the remaining papers to avoid per-token cost. Prompt version and lane are recorded per paper.")
H("3.5 Curation", 2)
P("Records are flattened, deduplicated on all fields, and checked against physical ranges (density 0.1–8 g/cm³; volume fraction ≤ 0.8; modulus 1–300,000 MPa; strength 0.1–3,000 MPa; wall-thickness ratio 0.3–1). Percent-not-fraction and GPa-not-MPa slips are flagged, not silently corrected. Specific modulus and strength are derived where density is present.")
H("3.6 Validation protocol", 2)
P(f"A stratified sample of {len(val)} primary records ({val.paper_id.nunique()} papers), balanced across matrix classes and across records with and without numeric properties, is provided with the source quote for each value. Two annotators, one a domain expert, will verify value, unit and row correctness; per-field precision, inter-annotator agreement (Cohen's κ) and an error taxonomy will be reported. This step is pending at the time of this draft.")
H("3.7 Benchmarks", 2)
P("ML. Targets: compressive modulus, compressive strength (compression tests only) and measured density. Features: matrix class, particle type, process route, test type, particle true density and diameter, wall-thickness ratio, volume and weight fraction, measured density (except when it is the target), strain rate and temperature. Only unflagged primary records are used. Ridge regression, random forest and gradient boosting are evaluated with 5-fold GroupKFold by paper so that no paper contributes to both training and test. Modulus and strength are modelled in log10.")
P("LLM. From the same pool, held-out records are sampled (40 strength, 40 density, 30 modulus). For each, the model receives the record's processing and test descriptors and is asked for a point estimate and a 90% credible interval, (a) zero-shot and (b) with the five nearest primary records from other papers (matched on matrix class, particle type and volume fraction) supplied as context. The model under test is Claude Opus 5 acting through Claude Code with file access restricted to the task text, so it cannot see the dataset or the source papers. Metrics: median and mean absolute percentage error, log-R², and empirical coverage of the 90% interval.")

# ---------------- 4 Results ----------------
H("4. Results")
H("4.1 What the literature contains", 2)
P(f"Of {n_proc} processed papers, {n_on} were confirmed on-topic; the {n_proc-n_on} rejected papers were keyword false positives (conservation reports using microballoon putty, an economics working paper, an antenna design, a radome FEA study) and produced no records. The on-topic papers yielded {len(df)} records: {prim} primary, {sec} secondary and {mod} model. The median paper contributes {int(df.groupby('paper_id').size().median())} records; theses contribute the most (up to {int(df.groupby('paper_id').size().max())}).")
T(["Quantity", "Value"], [
 ["Records with strength / modulus / measured density", f"{int(df.strength_mpa.notna().sum())} / {int(df.modulus_mpa.notna().sum())} / {int(df.measured_density_g_cc.notna().sum())}"],
 ["Records with particle volume fraction", f"{int(df.particle_volume_fraction.notna().sum())} (0–{df.particle_volume_fraction.max():.2f})"],
 ["Records at strain rate ≥ 100 s⁻¹", int((df.strain_rate_per_s >= 100).sum())],
 ["Matrix classes", ", ".join(f"{k} {v}" for k, v in df.matrix_class.value_counts().head(7).items())],
 ["Particle types", ", ".join(f"{k} {v}" for k, v in df.particle_type.value_counts().head(5).items())],
 ["Process routes", ", ".join(f"{k} {v}" for k, v in df.process_route.value_counts().head(5).items())],
 ["Test types", ", ".join(f"{k} {v}" for k, v in df.test_type.value_counts().head(6).items())],
 ["Records flagged by range checks", flagged],
], widths=[2.6, 3.9], caption="Table 1. Dataset composition.")
P(f"A central finding is how much of the literature is figure-only. Only {int(onpaper.sum())} of {len(onpaper)} on-topic papers state at least one numeric strength, modulus or density value in text or tables; the share rises from {pre:.0%} for papers before 2010 to {post:.0%} after (Figure 2). Records from figure-only papers carry composition and test conditions with null properties, which the extractor reports rather than digitising plots. Fire, tribological, acoustic and fracture-toughness results — present in several papers — have no schema field and are preserved in notes, indicating natural schema extensions.")
FIG(BENCH_DIR / "records_year_coverage.png", "Figure 2. Left: records per publication year. Right: fraction of records carrying strength, modulus and density, by matrix class.")
H("4.2 Extracted data reproduce known structure–property relations", 2)
P(f"Figure 3 plots density against compressive modulus and strength for unflagged primary compression records. Metal-matrix foams occupy the high-density, high-strength region (aluminium: {al.measured_density_g_cc.min():.2f}–{al.measured_density_g_cc.max():.2f} g/cm³, {al.strength_mpa.min():.0f}–{al.strength_mpa.max():.0f} MPa); polymer foams cluster at 0.3–0.95 g/cm³. Within aluminium foams, strength correlates with density (r = {r_al:.2f}, n = {len(al)}) despite spanning different matrices (Al99.5, AlSi12, A356, 7075), particle chemistries and strain rates; within a single study the correlation is near-perfect (r = 0.98 for the six quasi-static Al7075 compositions of one paper). The split-Hopkinson series in the dataset show the expected 2–3× strain-rate strengthening of Al-matrix foams.")
FIG(BENCH_DIR / "pilot_density_strength.png", "Figure 3. Density versus compressive modulus (left) and strength (right), unflagged primary records, coloured by matrix class.")
H("4.3 Cross-laboratory machine-learning baselines", 2)
T(["Target", "Model", "n rows / papers", "R² (log for E, σ)", "MAPE %"],
  [[m.target, m.model, f"{m.n} / {m.papers}", f"{m.r2_log:.2f}" if pd.notna(m.r2_log) else f"{m.r2:.2f}", f"{m.mape_pct:.0f}"] for _, m in ml.iterrows()],
  widths=[1.6, 1.2, 1.3, 1.3, 1.0], caption="Table 2. Paper-level GroupKFold results.")
best_d = ml[ml.target == "measured_density_g_cc"].r2.max(); best_e = ml[ml.target == "modulus_mpa"].r2_log.max(); best_s = ml[ml.target == "strength_mpa"].r2_log.max()
P(f"Density is predictable from processing descriptors across laboratories (random forest R² = {best_d:.2f}, MAPE ≈ 20%), as expected from the rule of mixtures. Modulus and strength are barely predictable: the best log-R² values are {best_e:.2f} and {best_s:.2f}, with MAPE around 90–110%. The rows mix metal and polymer matrices, quasi-static and dynamic loading, and heterogeneous particle grades, and most papers report a single matrix with one to six compositions, so the paper-level split forces extrapolation to unseen matrix/particle combinations. This is the quantitative statement of the data-scarcity problem: at ~140 usable strength rows from 23 papers, recipe-level features do not transfer. Data quality matters as much as quantity: before a curation rule flagged metal-matrix moduli reported in GPa under an MPa label (ISO 13314 'structural stiffness' in several papers), the modulus log-R² was −0.89; removing those 17 rows raised it to {best_e:.2f}.")
H("4.4 LLM prediction and calibration", 2)
if llm is not None and len(llm):
    T(["Target", "Condition", "n", "median APE %", "mean APE %", "log-R²", "90% coverage"],
      [[r.target, r.condition, int(r.n), f"{r.median_ape_pct:.0f}", f"{r.mape_pct:.0f}", f"{r.r2_log:.2f}", f"{r.coverage_90:.0%}"] for _, r in llm.iterrows()],
      widths=[1.5, 0.9, 0.5, 1.0, 1.0, 0.8, 1.0], caption="Table 3. LLM (Claude Opus 5) prediction on held-out records, zero-shot and with five retrieved records from other papers.")
    fig_p = BENCH_DIR / "llm_vs_ml.png"
    if fig_p.exists(): FIG(fig_p, "Figure 4. LLM predictions versus measured values (log axes), zero-shot and retrieval-augmented; dashed line is identity.")
    P("[Interpretation to be written once numbers are final: compare LLM zero-shot MAPE with the ML baselines in Table 2 on the same targets; report whether the 90% intervals are over- or under-confident; note where retrieval helps most (matrix classes with dense literature) and where it fails.]")
else:
    P("[LLM benchmark running — table and figure inserted automatically when data/benchmarks/llm_agent_summary.csv exists.]")

# ---------------- 5 Discussion ----------------
H("5. Discussion")
B([
 "Provenance is the enabling design choice. Because every value carries a quote and location, the expert audit reduces to checking quotes, and internal inconsistencies found by the extractor (table vs. text values, unit labels contradicting magnitudes, negative porosities from sign conventions) become a curated list rather than silent errors.",
 "Refusing to guess is measurable. Null properties in on-topic papers are a feature of the literature, not the extractor; the pre/post-2010 contrast (59% vs 86% of papers with numeric values) quantifies the return on a figure-digitisation stage.",
 "Secondary data must be tagged. Review tables and comparison rows contribute 16% of records; without the data-origin tag they would double-count originals and leak between train and test.",
 "Cross-laboratory prediction is the right target and it is hard. Density transfers; mechanical properties do not at ~10² rows. The dataset therefore serves less as a training set today than as a benchmark and a map of where measurements are missing.",
 "Cost. The entire pilot required ≈US$1 of API spend plus subagent compute; the full 1,329-paper corpus is estimated at US$400–500 via batch API, dominated by input tokens for long theses.",
])
H("6. Limitations")
B([
 f"Open-access bias: only {n_oa/len(rel):.0%} of relevant papers are open access; institutional access is required for the remainder and may shift the composition of the dataset toward journals with different reporting norms.",
 "Figure-only data are not captured; a plot-digitisation stage with uncertainty is future work.",
 "Field mapping strain: some papers report properties the schema lacks (fracture toughness, hardness, acoustic, fire), and some values were mapped to the nearest field with a note (e.g., ISO 13314 structural stiffness recorded as modulus).",
 "Validation is pending; extraction accuracy is currently supported only by spot checks and automated range flags.",
 "The LLM under test also performed extraction (different sessions, no file access during prediction); an independent model family should be added.",
])
H("7. Data and Code Availability")
P("Dataset (CSV/Parquet), per-paper extractions with evidence (JSONL), prompts, schema, curation and benchmark code: https://github.com/Harshelite2503/foamgpt (MIT). A Zenodo DOI will accompany the validated release.")
H("References")
B([
 "Jiang X. et al. Applications of natural language processing and large language models in materials discovery. npj Comput. Mater. 2025.",
 "Van M.H. et al. A survey of AI for materials science: foundation models, LLM agents, datasets, and tools. ACM Comput. Surv. 2026.",
 "Zimmermann Y. et al. 32 examples of LLM applications in materials science and chemistry. Mach. Learn.: Sci. Technol. 2025.",
 "Schilling-Wilhelmi M. et al. From text to insight: large language models for chemical data extraction. Chem. Soc. Rev. 2025.",
 "Zhang J. et al. MatSciBench: benchmarking the reasoning ability of LLMs in materials science. 2026.",
 "van Herck J. et al. Assessment of fine-tuned LLMs for real-world chemistry and material science applications. Chem. Sci. 2025.",
 "Gupta N., Rohatgi P.K. (eds). Metal Matrix Syntactic Foams. DEStech, 2014.",
 "Orbulov I.N., Dobránszky J. Producing metal matrix syntactic foams by pressure infiltration. Period. Polytech. Mech. Eng. 2008.",
 "Porfiri M., Gupta N. Effect of volume fraction and wall thickness on the elastic properties of hollow particle filled composites. Compos. B 2009.",
 "Beckwith C., Gupta N. Rheological measurement dataset of resin and composite mixtures for DLP additive manufacturing. Sci. Data 2026.",
])
out = ROOT / "FoamGPT_Draft_Paper.docx"; doc.save(out); (ROOT / "docs" / out.name).write_bytes(out.read_bytes())
(ROOT / "docs" / "paper_draft.md").write_text("\n".join(md)); print("saved", out)
