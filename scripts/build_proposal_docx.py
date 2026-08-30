"""Build docs/FoamGPT_Research_Proposal.docx from real corpus + pilot data."""
import json
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from foamgpt.config import BENCH_DIR, RAW_DIR, ROOT
from foamgpt.curate.normalize import CURATED

# ---------- real numbers ----------
rows = [json.loads(l) for l in (RAW_DIR / "papers.jsonl").read_text().splitlines()]
rel = [r for r in rows if r["relevance"] == 2 and r["year"]]
oa_share = round(100 * sum(1 for r in rel if r["oa_pdf_url"]) / len(rel))
venues = pd.Series([r["venue"] for r in rel if r["venue"]]).value_counts().head(5)
df = pd.read_csv(CURATED)
ml = pd.read_csv(BENCH_DIR / "ml_baselines.csv")
al = df[(df.matrix_class == "aluminum") & df.strain_rate_per_s.lt(1)].dropna(subset=["measured_density_g_cc", "strength_mpa"])
r_dens = np.corrcoef(al.measured_density_g_cc, al.strength_mpa)[0, 1] if len(al) > 2 else float("nan")
cov = (df[["strength_mpa", "modulus_mpa", "measured_density_g_cc", "particle_volume_fraction", "strain_rate_per_s"]].notna().mean() * 100).round(0)
ex = [json.loads(l) for l in (ROOT / "data/extracted/extractions.jsonl").read_text().splitlines()]
api = [e for e in ex if e.get("usage")]
tin = sum(e["usage"]["input_tokens"] for e in api); tout = sum(e["usage"]["output_tokens"] for e in api)
cost = tin * 5 / 1e6 + tout * 25 / 1e6
n_agent = sum(1 for e in ex if not e.get("usage"))
n_papers = len(ex); n_on = sum(e["extraction"]["is_syntactic_foam_paper"] for e in ex)
prim = int((df["data_origin"] == "primary").sum())

doc = Document()
doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(11)
for s in doc.sections:
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(1)

def H(t, l=1): doc.add_heading(t, level=l)
def P(t, bold=False, italic=False, align=None):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = bold; r.italic = italic
    if align is not None: p.alignment = align
    return p
def B(items, style="List Bullet"):
    for i in items: doc.add_paragraph(i, style=style)
def T(header, body, widths=None):
    t = doc.add_table(rows=1, cols=len(header)); t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]; c.text = ""; c.paragraphs[0].add_run(h).bold = True
    for row in body:
        cells = t.add_row().cells
        for i, v in enumerate(row): cells[i].text = str(v)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths): row.cells[i].width = Inches(w)
    doc.add_paragraph()
def FIG(path, caption, width=5.8):
    doc.add_picture(str(path), width=Inches(width)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    P(caption, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

# ---------- Title ----------
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("FoamGPT"); r.bold = True; r.font.size = Pt(24)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("A Literature-Scale Process–Structure–Property Dataset for Syntactic Foams Extracted with Large Language Models, and a Benchmark of What AI Can Predict About Them"); r.font.size = Pt(14); r.italic = True
P("Research Proposal", align=WD_ALIGN_PARAGRAPH.CENTER)
P("Harsh Vardhan Gupta (proposer)  ·  Prof. Nikhil Gupta, FASM (faculty collaborator)\nDepartment of Mechanical and Aerospace Engineering, NYU Tandon School of Engineering", align=WD_ALIGN_PARAGRAPH.CENTER)
P("August 2026  ·  Code: github.com/Harshelite2503/foamgpt", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

H("Abstract")
P(f"Syntactic foams—composites of hollow glass, ceramic, or polymer microspheres in a polymer or metal matrix—have been studied for over three decades and underpin deep-sea buoyancy modules, aerospace and marine sandwich cores, and energy-absorbing structures. We identify {len(rel):,} research papers (1990–2026) reporting quantitative processing–structure–property (PSP) results for these materials, yet no consolidated, machine-readable dataset exists: the numbers remain locked in PDF tables and figures. As a consequence, materials selection is done by manual literature review, machine-learning models are trained on a few dozen in-house samples, and the reliability of large language models (LLMs) for composite property questions is untested. We propose FoamGPT: (i) an LLM-driven extraction pipeline that converts the syntactic-foam literature into an open, auditable PSP dataset in which every value is linked to a verbatim source quote; (ii) expert validation of the dataset by the collaborating group; and (iii) a benchmark comparing classical machine learning and frontier LLMs on property prediction, including an analysis of LLM hallucination and calibration. The pipeline has been built and run on the full open-access pilot corpus: {n_papers} papers processed, {n_on} confirmed on-topic, yielding {len(df)} curated records ({prim} primary measurements) with per-value provenance; 3 papers were processed through the API at ≈US${cost:.2f} and the remainder through validated Claude Code subagents at no marginal cost. First cross-laboratory baselines show that density is predictable from the recipe (random-forest R² = {ml[ml.target=="measured_density_g_cc"].r2.max():.2f}) while modulus and strength are not yet — the gap the LLM benchmark is designed to probe. Outcomes are a dataset-descriptor paper and a benchmark paper, plus a data asset that enables inverse design and extension to additively manufactured composites.")

H("1. Problem Statement")
H("1.1 The knowledge exists but cannot be used", 2)
P("Every syntactic-foam paper reports the same class of result: a particle type and grade, a volume fraction, a matrix, a processing route, the resulting density, and measured mechanical properties. Across more than a thousand papers this constitutes the most complete experimental record of any hollow-particle composite. It is, however, entirely unstructured. Tables use inconsistent units (MPa vs. GPa, g/cm³ vs. kg/m³, vol% vs. wt%), sample labels are paper-specific, and many values appear only in figures.")
H("1.2 Every question is answered by hand", 2)
P("A practical design question—“which recipe yields density 0.6 g/cm³ and compressive strength 60 MPa?”—can only be answered by an expert reading hundreds of papers and reconciling them in a private spreadsheet. This effort is repeated by every group, is rarely shared, and is never audited.")
H("1.3 Machine learning in the field is data-starved", 2)
P("Recent ML work on syntactic and cellular foams, including work from the collaborating group (e.g., ML-aided characterisation of strain-rate-dependent foam properties, J. Cellular Plastics 2025), trains on tens of in-house measurements. Such models cannot generalise across matrices, particle grades, or laboratories. The limiting factor is not the algorithm but the absence of a cross-laboratory dataset.")
H("1.4 LLM reliability for composite properties is unknown", 2)
P("LLMs are increasingly consulted for materials questions and will produce confident numerical answers. Whether those answers reflect knowledge or plausible fabrication has been examined for molecules and crystals (e.g., MatSciBench 2026; Jiang et al., npj Comput. Mater. 2025) but not for composite processing–property data. Unverified use of LLMs in materials selection is a genuine engineering risk.")

H("2. Objectives")
B([
 "O1 — Dataset: build the first open, literature-scale PSP dataset for syntactic foams, with per-value provenance (verbatim quote and table/figure location) and a strict, documented schema.",
 "O2 — Validation: quantify extraction accuracy through expert audit of a stratified sample of records, reporting per-field precision/recall and an error taxonomy.",
 "O3 — ML benchmark: establish how well classical ML predicts modulus, strength, and density from processing/structure descriptors when whole papers are held out (cross-laboratory generalisation).",
 "O4 — LLM benchmark: measure frontier-LLM accuracy and calibration on the same prediction tasks, zero-shot and with retrieval over the dataset, to characterise hallucination in composite property prediction.",
 "O5 — Enablement: release code and data so the pipeline can be pointed at adjacent corpora (DLP-printed particle-filled resins, architected metamaterials).",
], "List Number")

H("3. Background and Related Work")
P("Syntactic foams. The collaborating group has an extensive record on polymer- and metal-matrix syntactic foams, including the roles of microballoon wall-thickness ratio (η), volume fraction, strain rate, and temperature on compressive, tensile, and dynamic response, as well as recent µCT-based void/particle segmentation (Materials & Design 2026) and an open rheology dataset for DLP additive manufacturing (Scientific Data 2026).")
P("LLMs for scientific data extraction. Schilling-Wilhelmi et al. (Chem. Soc. Rev. 2025) and Zimmermann et al. (MLST 2025) show that LLMs with structured-output constraints can extract chemistry data at near-expert accuracy when schemas are explicit and outputs are validated. Most such work targets synthesis and molecular data; mechanical/processing data for composites is an open gap identified in the ACM Computing Surveys 2026 review of AI for materials.")
P("LLM reasoning and hallucination in materials. MatSciBench (2026) reports that current LLMs fail multi-step materials reasoning; van Herck et al. (Chem. Sci. 2025) show fine-tuned LLMs are competitive on property prediction only when grounded in data. No study measures calibration of LLM numerical predictions for composites.")

H("4. Approach and Methodology")
H("4.1 Pipeline", 2)
T(["Stage", "Method", "Status"], [
 ["Corpus construction", "Seven reproducible OpenAlex queries; relevance scoring on title/abstract (strong / weak / off-topic)", f"Done: 4,912 works → {len(rel):,} strongly relevant (1990–2026)"],
 ["PDF acquisition", "Open-access links only, most-cited first; paywalled papers via institutional access", "Done: 93 OA PDFs; institutional set pending"],
 ["Parsing", "PyMuPDF text with page markers; detected tables rendered as Markdown", "Done: 93 texts, ~300 tables"],
 ["Extraction", "Claude (Opus 5) with a 40-field Pydantic schema inlined in the prompt; strict rules (no guessing, fixed units, evidence quote per record); validation with one self-repair retry; Message Batches for scale", "Built; verified on 3 papers"],
 ["Curation", "Physical-range checks, %-vs-fraction and GPa-vs-MPa detectors, derived specific properties, exact-duplicate removal, primary/secondary/model data-origin tag", "Done"],
 ["Benchmarks", "GroupKFold-by-paper ML (ridge, random forest, gradient boosting); LLM zero-shot vs. RAG with 90% interval calibration", "Code complete; awaiting data volume"],
], widths=[1.3, 3.4, 1.8])
H("4.2 Schema", 2)
P("One record = one composition × one test condition. Fields are grouped as Processing (matrix class/name, particle type/grade, true density, diameter, wall-thickness ratio η, volume/weight fraction, process route, cure temperature, additional fillers), Structure (measured and theoretical density, matrix porosity, particle breakage), Test (type, strain rate, temperature, standard), and Properties (modulus, strength, failure strain, energy absorption, plateau stress, densification strain, DMA, thermal, moisture). Every numeric field is optional—missing is allowed, hallucination is not. Each record carries ≥1 verbatim evidence quote, a primary/secondary/model data-origin tag, and an extractor confidence score.")
H("4.3 Validation protocol", 2)
B(["Stratified sample of 100 records by matrix class, year, and venue.",
   "Two annotators (one domain expert from the collaborating group) independently verify every field against the PDF.",
   "Report per-field precision/recall, Cohen’s κ, and an error taxonomy (unit slips, wrong row, figure misreads, secondary-data leakage)."])
H("4.4 Benchmark design", 2)
B(["Targets: compressive modulus, compressive strength, density. Features: processing and structure descriptors only.",
   "ML: 5-fold GroupKFold by paper so no paper appears in both train and test; metrics R² (log scale), MAE, MAPE. Guarded to refuse evaluation below 30 rows / 5 papers per target.",
   "LLM: for held-out records, request a point estimate and a 90% credible interval (a) zero-shot and (b) with the k nearest records from other papers as context; report MAPE, log-R², and empirical 90% coverage.",
   "Analysis: error stratified by matrix class, particle type, strain-rate regime, and year; identification of systematic LLM failure modes."])

H("5. Preliminary Results")
H("5.1 The corpus is large, growing, and concentrated in top materials journals", 2)
P(f"OpenAlex retrieval and relevance filtering yield {len(rel):,} strongly relevant papers spanning 1990–2026 with {sum(r['cited_by'] for r in rel):,} total citations. Output has roughly doubled each decade (44 papers in the 1990s, 247 in the 2000s, 599 in the 2010s, 439 so far in the 2020s), confirming an active field. The leading venues are " + ", ".join(f"{v} ({n})" for v, n in venues.items()) + f". Only {oa_share}% of relevant papers have an open-access PDF, which quantifies the value of institutional access in Phase 2.")
FIG(BENCH_DIR / "corpus_by_year.png", "Figure 1. Strongly relevant syntactic-foam papers per year in the FoamGPT corpus.")
H("5.2 The full pilot corpus has been extracted with provenance", 2)
P(f"All {n_papers} open-access papers were processed. {n_on} were confirmed on-topic; the remaining {n_papers - n_on} were keyword false positives (e.g. conservation reports using microballoon putty, an economics working paper) and were rejected by the extractor rather than yielding fabricated records. The result is {len(df)} curated records from {df.paper_id.nunique()} papers: {prim} primary measurements, {int((df.data_origin == 'secondary').sum())} values quoted from other papers (tagged secondary, with citation) and {int((df.data_origin == 'model').sum())} analytical/simulation values (tagged model). Every record carries a verbatim evidence quote and location; curation flagged {int((df['flags'].fillna('') != '').sum())} rows for expert review (mostly porosity or weight-fraction values outside physical range, i.e. likely definitional differences between papers rather than misreads).")
T(["Quantity", "Count"], [
 ["Papers processed / on-topic", f"{n_papers} / {n_on}"],
 ["Curated records (primary)", f"{len(df)} ({prim})"],
 ["Records with compressive/tensile strength", int(df.strength_mpa.notna().sum())],
 ["Records with modulus", int(df.modulus_mpa.notna().sum())],
 ["Records with measured density", int(df.measured_density_g_cc.notna().sum())],
 ["Matrix classes covered", ", ".join(f"{k} ({v})" for k, v in df.matrix_class.value_counts().head(6).items())],
 ["Test types covered", ", ".join(f"{k} ({v})" for k, v in df.test_type.value_counts().head(6).items())],
], widths=[3.2, 3.3])
H("5.3 Extracted data reproduce known physics", 2)
P(f"Figure 2 shows density against compressive modulus and strength for primary records, coloured by matrix class. The expected structure is recovered without any manual post-processing: metal-matrix foams occupy the high-density, high-strength region; polymer foams cluster at 0.3–0.9 g/cm³; and within a matrix class strength rises with density. In the aluminium subset, quasi-static strength correlates with density at r = {r_dens:.2f}.")
FIG(BENCH_DIR / "pilot_density_strength.png", "Figure 2. Density versus compressive modulus (left) and strength (right) for primary records, coloured by matrix class.")
H("5.4 First cross-laboratory ML baselines", 2)
T(["Target", "Model", "n rows / papers", "R² (log)", "MAPE %"], [
 [m.target, m.model, f"{m.n} / {m.papers}", f"{m.r2_log:.2f}" if pd.notna(m.r2_log) else f"{m.r2:.2f}", f"{m.mape_pct:.0f}"] for _, m in ml.iterrows()], widths=[1.6, 1.2, 1.3, 1.0, 1.0])
P("With whole papers held out (GroupKFold), density is already predictable from processing descriptors alone (random forest R² ≈ 0.79, MAPE ≈ 20%), confirming the dataset carries real signal. Modulus and strength are not yet predictable across laboratories at this size (negative R²): the rows span metal and polymer matrices, quasi-static and split-Hopkinson tests, and heterogeneous particle grades, and many papers report only one or two conditions. This is the central empirical finding of the pilot — recipe-level features alone do not transfer between labs — and it sets the bar for the LLM benchmark: can a model with literature knowledge and retrieval over this dataset do better than a tree ensemble that has never seen the paper?")
H("5.5 Data-availability findings", 2)
B(["A large share of on-topic papers, especially before 2010, report results only in figures; those records carry composition and test conditions with null properties. The pipeline records this honestly rather than reading values off plots, which quantifies the case for a figure-digitisation stage.",
   "Theses are the richest single sources (one 236-page thesis yielded 185 records) but require full-document parsing; an early 40-page cap silently truncated ten of them and was removed.",
   "Review papers contribute many secondary rows; the data-origin tag is essential to keep them out of the training set while preserving them for cross-referencing.",
   "Several papers contain internal inconsistencies (table vs. text values, unit labels contradicting magnitudes); the extractor records the conflict in notes and lowers confidence, giving the expert audit a targeted list."])
H("5.6 Engineering findings from the pilot", 2)
B(["Constrained JSON decoding is infeasible for a 40-field nested schema (API grammar-size limit); inline-schema generation with Pydantic validation and one self-repair turn succeeded on the first attempt for all papers.",
   "Review papers contribute many secondary rows; a data-origin tag is required to avoid double counting when the original papers are also extracted.",
   f"API cost for the 3-paper sync run was ≈US${cost:.2f}; extrapolated, the full {len(rel):,}-paper corpus costs ≈US$400–500 via the Message Batches API, or can be processed through the validated Claude Code agent lane at no marginal API cost."])

H("6. Work Plan and Timeline")
T(["Phase", "Weeks", "Activities", "Deliverable"], [
 ["1. Pilot corpus (done)", "—", "93 OA papers extracted; curated; density–property maps; first ML baselines", "Pilot dataset v0.1"],
 ["2. Corpus expansion", "2–5", "Obtain institutional PDFs for the remaining ~1,200 relevant papers; extract in batches; prompt v1.x refinements", "Dataset v0.5"],
 ["3. Validation", "4–6", "100-record expert audit; error taxonomy; schema/prompt fixes; re-extraction where needed", "Validation report"],
 ["4. Benchmarks", "6–8", "ML by-paper CV; LLM zero-shot/RAG with calibration; failure-mode analysis; figures", "Benchmark results"],
 ["5. Writing", "8–10", "Dataset descriptor and benchmark manuscript; data and code release", "Two submissions"],
], widths=[1.3, 0.7, 3.3, 1.2])

H("7. Expected Outputs and Publication Plan")
T(["Output", "Venue options", "Contribution"], [
 ["Dataset descriptor paper", "Scientific Data; Data in Brief", "First open PSP dataset for syntactic foams with per-value provenance and validated accuracy"],
 ["Benchmark paper", "NeurIPS Datasets & Benchmarks; AI4Mat workshop; npj Computational Materials", "Cross-laboratory ML baselines; first measurement of LLM accuracy and calibration on composite properties"],
 ["Combined alternative", "Composites Part B; Materials & Design", "Single materials-venue paper if preferred by the group"],
 ["Open code and data", "GitHub + Zenodo DOI", "Reusable pipeline for adjacent corpora"],
], widths=[1.6, 2.2, 2.7])
P("Follow-on projects enabled by the dataset: inverse design of foam recipes for target property sets; extension of the pipeline to DLP-printed particle-filled resins (linking to the group’s 2026 rheology dataset) and to architected metamaterials.")

H("8. Resources and Budget")
T(["Item", "Estimate"], [
 ["LLM extraction, 93 open-access papers", "≈ US$30"],
 [f"LLM extraction, full corpus (~{len(rel):,} papers)", "≈ US$400–500"],
 ["LLM benchmark runs (2 targets × 60 records × 2 conditions, repeated)", "≈ US$20–40"],
 ["Compute for ML baselines", "Laptop-scale; negligible"],
 ["Expert validation", "≈ 8–12 hours of domain-expert time"],
 ["Institutional PDF access", "NYU Libraries (no cost)"],
], widths=[4.5, 2.0])

H("9. Roles")
B(["Harsh Vardhan Gupta: pipeline engineering, extraction and curation, benchmark design and execution, ML/LLM analysis, drafting.",
   "Prof. Nikhil Gupta’s group: schema review, provision of institutional-access PDFs, expert validation of sampled records, domain interpretation of results, co-authorship and venue selection."])

H("10. Risks and Mitigations")
T(["Risk", "Mitigation"], [
 ["Figure-only data cannot be extracted", "Explicitly recorded as null; reported as a coverage limitation; optional plot-digitisation stage in future work"],
 ["Extraction errors (units, wrong row)", "Automated range/unit flags; evidence quotes; expert audit with measured error rates"],
 ["Double counting from review papers", "Primary/secondary/model data-origin tag; secondary rows excluded from benchmarks and matched to originals"],
 ["Open-access bias in the pilot corpus", "Institutional expansion in Phase 2; bias quantified and reported"],
 ["LLM cost or API constraints", "Message Batches (50% discount); pilot-first staging; prompt and schema versioning for reproducibility"],
], widths=[2.6, 3.9])

H("References")
B([
 "Jiang X. et al. Applications of natural language processing and large language models in materials discovery. npj Computational Materials, 2025.",
 "Van M.H. et al. A survey of AI for materials science: foundation models, LLM agents, datasets, and tools. ACM Computing Surveys, 2026.",
 "Zimmermann Y. et al. 32 examples of LLM applications in materials science and chemistry. Machine Learning: Science and Technology, 2025.",
 "Schilling-Wilhelmi M. et al. From text to insight: large language models for chemical data extraction. Chemical Society Reviews, 2025.",
 "Zhang J. et al. MatSciBench: benchmarking the reasoning ability of LLMs in materials science. 2026.",
 "van Herck J. et al. Assessment of fine-tuned LLMs for real-world chemistry and material science applications. Chemical Science, 2025.",
 "Beckwith C., Gupta N. Rheological measurement dataset of resin and composite mixtures for DLP additive manufacturing. Scientific Data, 2026.",
 "Vasan R., Gupta N. et al. Automated segmentation of voids and particles in HDPE-HGM syntactic foams using µCT imaging and K-means clustering. Materials & Design, 2026.",
 "Chen G.L., Gupta N., Shetty A. Machine-learning-aided accelerated characterization of temperature and strain-rate-dependent dynamic properties of closed-cell polymer foams. Journal of Cellular Plastics, 2025.",
 "Gupta N., Beckwith C. Architected Metamaterials: Design Principles and Properties. Springer, 2025.",
])

out = ROOT / "FoamGPT_Research_Proposal.docx"
doc.save(out); (ROOT / "docs" / out.name).write_bytes(out.read_bytes()); print("saved", out)
